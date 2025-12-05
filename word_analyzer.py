import os
import re
import sys
import tempfile
import datetime
import webbrowser
from pathlib import Path
from flask import Flask, request, render_template_string, send_from_directory, redirect, url_for
from collections import defaultdict
from dataclasses import dataclass
from typing import List, Tuple, Dict, Any
import pythoncom
from flask import Flask, request, render_template_string, send_file, redirect, url_for, jsonify
import threading

progress_data = {}
# --- Detect optional libraries safely ---
try:
    import win32com.client
    HAS_WIN32COM = True
except Exception:
    HAS_WIN32COM = False

try:
    from docx import Document
    HAS_DOCX = True
except Exception:
    HAS_DOCX = False

# ---------------------------------------------------------------------
# 1️⃣ BuildDetailedSummaryTable (Python conversion)
# ---------------------------------------------------------------------
def build_detailed_summary_table(
    dict_types: dict,
    figure_count: int,
    table_count: int,
    footnote_count: int,
    endnote_count: int,
    fmt_content: str,
    spec_content: str,
    comment_content: str
) -> str:
    """Equivalent to VBA BuildDetailedSummaryTable()"""

    def count_items(section_html: str, token: str) -> int:
        return section_html.lower().count(token.lower())

    def build_progress_row(title: str, cap_cnt: int, cit_cnt: int, miss_cap: int, miss_cit: int) -> str:
        total = max(cap_cnt, cit_cnt)
        complete_pct = round(((total - miss_cap - miss_cit) / total * 100), 1) if total else 0
        html = f"""
        <tr>
          <td><strong>{title}</strong></td>
          <td>{total}</td>
          <td>
            <div style='display:flex;align-items:center;gap:10px;'>
              <div style='width:100px;height:20px;background:#f0f0f0;border-radius:10px;overflow:hidden;'>
                <div style='width:{complete_pct}%;height:100%;background:linear-gradient(90deg,#27ae60,#2ecc71);'></div>
              </div>
              <span style='font-size:12px;color:#27ae60;'>{complete_pct}% Complete</span>
            </div>
          </td>
          <td>
            <i class='fas fa-check-circle' style='color:#27ae60;'></i> {cit_cnt} citation(s)<br>
            {'<span style="color:#e74c3c;"><i class="fas fa-times-circle"></i> Missing ' + str(miss_cap) + ' caption(s)</span>' if miss_cap else ''}
          </td>
          <td>
            <i class='fas fa-check-circle' style='color:#27ae60;'></i> {cap_cnt} caption(s)<br>
            {'<span style="color:#f39c12;"><i class="fas fa-exclamation-triangle"></i> Missing ' + str(miss_cit) + ' citation(s)</span>' if miss_cit else ''}
          </td>
          <td>{"Add missing items" if miss_cap or miss_cit else "No action required"}</td>
        </tr>
        """
        return html

    def build_critical_issues_block(fig_miss_cap, fig_miss_cit, tab_miss_cap, tab_miss_cit, fmt_count):
        html = """
        <div style='background:#fff3cd;border:1px solid #ffeaa7;border-radius:10px;padding:20px;margin-top:20px;'>
          <h3 style='color:#856404;margin-bottom:15px;'><i class='fas fa-exclamation-triangle'></i> Critical Issues Requiring Attention</h3>
          <ul style='margin:0;padding-left:20px;color:#856404;'>
        """
        if (fig_miss_cit + tab_miss_cit) > 0:
            html += f"<li><strong>{fig_miss_cit + tab_miss_cit} Missing Citations:</strong> Check missing citations in Citations tab</li>"
        if (fig_miss_cap + tab_miss_cap) > 0:
            html += f"<li><strong>{fig_miss_cap + tab_miss_cap} Missing Captions:</strong> Check missing captions in Citations tab</li>"
        if fmt_count > 0:
            html += f"<li><strong>{fmt_count} Formatting Issues:</strong> See Formatting tab</li>"
        html += "</ul></div>"
        return html

    # --- counts ---
    fmt_count = count_items(fmt_content, "<tr")
    spec_count = count_items(spec_content, "<tr")
    comment_count_val = count_items(comment_content, "<tr")

    # --- figure/table stats ---
    fig_cap = fig_cit = fig_miss_cap = fig_miss_cit = 0
    tab_cap = tab_cit = tab_miss_cap = tab_miss_cit = 0

    def normalize_ref(ref: str) -> str:
        return ref.replace("-", ".").strip().lower()

    for type_key in dict_types.keys():
        if type_key == "Figure":
            fig_cap = len(dict_types[type_key]["Caption"])
            fig_cit = len(dict_types[type_key]["Citation"])
            for k in dict_types[type_key]["Citation"]:
                norm = normalize_ref(k)
                if not any(normalize_ref(x) == norm for x in dict_types[type_key]["Caption"]):
                    fig_miss_cap += 1
            for k in dict_types[type_key]["Caption"]:
                norm = normalize_ref(k)
                if not any(normalize_ref(x) == norm for x in dict_types[type_key]["Citation"]):
                    fig_miss_cit += 1
        elif type_key == "Table":
            tab_cap = len(dict_types[type_key]["Caption"])
            tab_cit = len(dict_types[type_key]["Citation"])
            for k in dict_types[type_key]["Citation"]:
                norm = normalize_ref(k)
                if not any(normalize_ref(x) == norm for x in dict_types[type_key]["Caption"]):
                    tab_miss_cap += 1
            for k in dict_types[type_key]["Caption"]:
                norm = normalize_ref(k)
                if not any(normalize_ref(x) == norm for x in dict_types[type_key]["Citation"]):
                    tab_miss_cit += 1

    # --- build HTML summary ---
    html = """
    <div class='header'>
      <div class='section-title'><i class='fas fa-chart-pie'></i> Analysis Summary</div>
      <table style='margin-bottom:20px;width:100%;border-collapse:collapse;'>
        <thead>
          <tr>
            <th>Element Type</th>
            <th>Total Found</th>
            <th>Status Overview</th>
            <th>Citations Status</th>
            <th>Captions Status</th>
            <th>Action Required</th>
          </tr>
        </thead><tbody>
    """

    html += build_progress_row("Figures", fig_cap, fig_cit, fig_miss_cap, fig_miss_cit)
    html += build_progress_row("Tables", tab_cap, tab_cit, tab_miss_cap, tab_miss_cit)

    # add special chars, formatting, comments, notes
    html += f"""
    <tr><td><strong>Special Characters</strong></td><td>{spec_count}</td>
        <td colspan='3'><a href='javascript:void(0);' onclick="showTab('special-chars');"
        style='color:#667eea;text-decoration:underline;'>Review multilingual symbols</a></td>
        <td>Review unusual characters</td></tr>

    <tr><td><strong>Formatting Issues</strong></td><td>{fmt_count}</td>
        <td colspan='3'><a href='javascript:void(0);' onclick="showTab('formatting');"
        style='color:#f39c12;text-decoration:underline;'>View formatting issues</a></td>
        <td>Review formatting anomalies</td></tr>

    <tr><td><strong>Comments</strong></td><td>{comment_count_val}</td>
        <td colspan='3'><a href='javascript:void(0);' onclick="showTab('comments');"
        style='color:#3498db;text-decoration:underline;'>Review editor comments</a></td>
        <td>Review highlighted feedback</td></tr>

    <tr><td><strong>Notes</strong></td><td>{footnote_count + endnote_count}</td>
        <td colspan='3'><a href='javascript:void(0);' onclick="showTab('media');"
        style='color:#27ae60;text-decoration:underline;'>{footnote_count} Footnotes, {endnote_count} Endnotes</a></td>
        <td>No action required</td></tr>
    """

    if figure_count > 0:
        html += f"""
        <tr><td><strong>Images</strong></td><td>{figure_count}</td>
        <td colspan='3'><a href='javascript:void(0);' onclick="showTab('media');"
        style='color:#27ae60;text-decoration:underline;'><i class='fas fa-check-circle'></i> {figure_count} image(s) detected</a></td>
        <td>No action required</td></tr>
        """
    else:
        html += """
        <tr><td><strong>Images</strong></td><td>0</td>
        <td colspan='3'><span style='color:#e67e22;'><i class='fas fa-exclamation-triangle'></i> No images detected</span></td>
        <td>Check for missing image elements</td></tr>
        """

    html += "</tbody></table>"
    html += build_critical_issues_block(fig_miss_cap, fig_miss_cit, tab_miss_cap, tab_miss_cit, fmt_count)
    html += "</div>"

    return html
# ------------------------------
# Citation / Caption Analyzer
# ------------------------------
@dataclass
class CitationItem:
    item_id: str
    page_no: int
    is_caption: bool


class CitationAnalyzer:
    def __init__(self):
        self.supported_types = ["Figure", "Table", "Box", "Exhibit", "Appendix", "Case Study"]
        self.regex_patterns = self._setup_regex_patterns()

    def _setup_regex_patterns(self) -> Dict[str, re.Pattern]:
        patterns = {}
        patterns['single'] = re.compile(
            r'(?:\(|\b)(Figure|Fig\.?|Table|Tab\.?|Box|Exhibit|Appendix|Case\s+Study)\.?\s*([0-9]+(?:[.\-][0-9]+)*)([A-Za-z]?)(?:\)|\b)',
            re.IGNORECASE
        )
        patterns['range'] = re.compile(
            r'(?:\(|\b)(Figures?|Figs?\.?|Tables?|Tabs?\.?|Boxes?|Exhibits?|Appendices?|Case\s+Studies?)\.?\s+([0-9]+(?:[\.\-][0-9]+)+)([A-Za-z]?)\s*(?:to|through|–|—|-)\s*([0-9]+(?:[\.\-][0-9]+)*)([A-Za-z]?)(?:\)|\b)',
            re.IGNORECASE
        )
        patterns['and'] = re.compile(
            r'(?:\(|\b)(Figures?|Figs?\.?|Tables?|Tabs?\.?|Boxes?|Exhibits?|Appendices?|Case\s+Studies?)\.?\s+([0-9]+(?:[\.\-][0-9]+)+)([A-Za-z]?)\s+(?:and|&)\s*([0-9]+(?:[\.\-][0-9]+)*)([A-Za-z]?)(?:\)|\b)',
            re.IGNORECASE
        )
        return patterns

    def normalize_for_regex(self, text: str) -> str:
        text = text.replace('\u2013', '-').replace('\u2014', '-').replace('\xa0', ' ')
        return text

    def normalize_type(self, label: str) -> str:
        if not label:
            return "Figure"
        lbl = label.lower()
        if lbl.startswith('fig'):
            return "Figure"
        if lbl.startswith('tab'):
            return "Table"
        if lbl.startswith('box'):
            return "Box"
        if lbl.startswith('exhibit'):
            return "Exhibit"
        if lbl.startswith('appendix'):
            return "Appendix"
        if lbl.startswith('case'):
            return "Case Study"
        return "Figure"

    def normalize_fig_number(self, fig_ref: str) -> str:
        if not fig_ref:
            return ""
        fig_ref = fig_ref.strip()
        fig_ref = fig_ref.replace('--', '-').replace('\u2013', '-').replace('\u2014', '-')
        for ch in ['[', ']', '°']:
            fig_ref = fig_ref.replace(ch, '')
        m = re.search(r'([0-9]+(?:[.\-][0-9]+)*)([A-Za-z]?)', fig_ref)
        if m:
            base = m.group(1).replace('-', '.')
            suffix = m.group(2)
            if base.endswith('.'):
                base = base[:-1]
            return base + suffix
        return fig_ref

    def is_caption_paragraph(self, text: str) -> bool:
        t = self.normalize_for_regex(text.strip()).lower()
        if not t:
            return False
        if len(t.splitlines()) > 7:
            return False
        for prefix in ['figure', 'fig.', 'table', 'tab.', 'box', 'exhibit', 'appendix', 'case study']:
            if t.startswith(prefix):
                return True
        return False

    def analyze_document_citations(self, document_content: List[Tuple[str, int, bool]]) -> Dict[str, Any]:
        dict_types = {t: {"Caption": {}, "Citation": {}, "CaptionPage": {}, "CitationPage": {}} for t in self.supported_types}

        for text, page_no, is_caption in document_content:
            txt = self.normalize_for_regex(text)

            for m in self.regex_patterns['range'].finditer(txt):
                label = self.normalize_type(m.group(1))
                start_num = self.normalize_fig_number(m.group(2))
                end_num = self.normalize_fig_number(m.group(4))
                try:
                    sp = start_num.split('.')
                    ep = end_num.split('.')
                    if int(sp[0]) == int(ep[0]) and len(sp) > 1 and len(ep) > 1:
                        start_minor = int(sp[1])
                        end_minor = int(ep[1])
                        for n in range(start_minor, end_minor + 1):
                            item_id = f"{label} {sp[0]}.{n}"
                            self._store(dict_types, label, item_id, page_no, is_caption)
                    else:
                        self._store(dict_types, label, f"{label} {start_num}", page_no, is_caption)
                        self._store(dict_types, label, f"{label} {end_num}", page_no, is_caption)
                except Exception:
                    self._store(dict_types, label, f"{label} {start_num}", page_no, is_caption)
                    self._store(dict_types, label, f"{label} {end_num}", page_no, is_caption)

            for m in self.regex_patterns['and'].finditer(txt):
                label = self.normalize_type(m.group(1))
                first_num = self.normalize_fig_number(m.group(2))
                second_num = self.normalize_fig_number(m.group(4))
                self._store(dict_types, label, f"{label} {first_num}", page_no, is_caption)
                self._store(dict_types, label, f"{label} {second_num}", page_no, is_caption)

            for m in self.regex_patterns['single'].finditer(txt):
                label = self.normalize_type(m.group(1))
                main_no = m.group(2)
                suffix = m.group(3) or ""
                item_id = f"{label} {self.normalize_fig_number(main_no + suffix)}"
                self._store(dict_types, label, item_id, page_no, is_caption)

        return dict_types

    def _store(self, dict_types, label, item_id, page_no, is_caption):
        tdict = dict_types.get(label)
        if tdict is None:
            return
        if is_caption:
            if item_id not in tdict['Caption']:
                tdict['Caption'][item_id] = True
                tdict['CaptionPage'][item_id] = page_no
        else:
            if item_id not in tdict['Citation']:
                tdict['Citation'][item_id] = True
                tdict['CitationPage'][item_id] = page_no

    def build_citation_tables_html(self, dict_types: Dict, doc_name: str) -> str:
        html = "<div class='citation-analysis'>"
        html += self._build_summary_table(dict_types)
        html += self._build_table("Citations Found", dict_types, "Citation", doc_name)
        html += self._build_table("Captions Found", dict_types, "Caption", doc_name)
        html += self._build_missing_table("Missing Captions", dict_types, True, doc_name)
        html += self._build_missing_table("Missing Citations", dict_types, False, doc_name)
        html += "</div>"
        return html

    def _build_summary_table(self, dict_types):
        h = "<h3>Summary Overview</h3><table class='summary-table'><thead><tr><th>Type</th><th>Captions</th><th>Citations</th><th>Missing Captions</th><th>Missing Citations</th></tr></thead><tbody>"
        for type_key in self.supported_types:
            cap_cnt = len(dict_types[type_key]["Caption"])
            cit_cnt = len(dict_types[type_key]["Citation"])
            miss_cap_cnt = sum(
                1 for cit_key in dict_types[type_key]["Citation"].keys()
                if not any(self.normalize_fig_number(cap_key) == self.normalize_fig_number(cit_key)
                           for cap_key in dict_types[type_key]["Caption"].keys())
            )
            miss_cit_cnt = sum(
                1 for cap_key in dict_types[type_key]["Caption"].keys()
                if not any(self.normalize_fig_number(cit_key) == self.normalize_fig_number(cap_key)
                           for cit_key in dict_types[type_key]["Citation"].keys())
            )
            if cap_cnt > 0 or cit_cnt > 0:
                h += f"<tr><td><strong>{type_key}</strong></td><td>{cap_cnt}</td><td>{cit_cnt}</td><td>{miss_cap_cnt}</td><td>{miss_cit_cnt}</td></tr>"
        h += "</tbody></table>"
        return h

    def _build_table(self, title, dict_types, dict_key, doc_name):
        h = f"<h3>{title}</h3><table id='{title.replace(' ', '').lower()}Table'><thead><tr><th>Document</th><th>Type</th><th>Item</th><th>Page</th></tr></thead><tbody>"
        count = 0
        for type_key in self.supported_types:
            for item_key in sorted(dict_types[type_key][dict_key].keys()):
                page_no = dict_types[type_key].get(dict_key + "Page", {}).get(item_key, "N/A")
                h += f"<tr><td>{doc_name}</td><td>{type_key}</td><td>{item_key}</td><td>{page_no}</td></tr>"
                count += 1
        if count == 0:
            h += "<tr><td colspan='4'>No items found</td></tr>"
        h += "</tbody></table>"
        return h

    def _build_missing_table(self, title, dict_types, missing_cap, doc_name):
        h = f"<h3>{title}</h3><table id='{title.replace(' ', '').lower()}Table'><thead><tr><th>Document</th><th>Type</th><th>Item</th><th>Page</th></tr></thead><tbody>"
        count = 0
        for type_key in self.supported_types:
            if missing_cap:
                for cit_key in dict_types[type_key]["Citation"].keys():
                    if not any(self.normalize_fig_number(cap_key) == self.normalize_fig_number(cit_key)
                               for cap_key in dict_types[type_key]["Caption"].keys()):
                        page_no = dict_types[type_key]["CitationPage"].get(cit_key, "N/A")
                        h += f"<tr><td>{doc_name}</td><td>{type_key}</td><td>{cit_key}</td><td>{page_no}</td></tr>"
                        count += 1
            else:
                for cap_key in dict_types[type_key]["Caption"].keys():
                    if not any(self.normalize_fig_number(cit_key) == self.normalize_fig_number(cap_key)
                               for cit_key in dict_types[type_key]["Citation"].keys()):
                        page_no = dict_types[type_key]["CaptionPage"].get(cap_key, "N/A")
                        h += f"<tr><td>{doc_name}</td><td>{type_key}</td><td>{cap_key}</td><td>{page_no}</td></tr>"
                        count += 1
        if count == 0:
            h += "<tr><td colspan='4'>All items matched</td></tr>"
        h += "</tbody></table>"
        return h

    def remove_tags_keep_formatting_docx(self, path):
        pass


def remove_tags_keep_formatting_docx(self, doc_path):
    """
    Full cleanup for Word file:
    ✔ Removes all <tags>
    ✔ Removes leading spaces in paragraphs
    ✔ Converts double/multiple spaces → single space
    ✔ Removes blank lines
    ✔ Works on ALL Word story ranges (headers, footers, textboxes, comments, footnotes)
    """

    import pythoncom, win32com.client, re
    pythoncom.CoInitialize()

    # === Patterns ===
    tag_patterns = [
        r"<[^ >]+>",
        r"</[^ >]+>",
        r"</([0-9A-Za-z]+)>",
        r"<([0-9A-Za-z]+)>",
        r"<([0-9A-Za-z_-]+)[^ >]*>",
        r"</([0-9A-Za-z_-]+)[^ >]*>",
        r"</([0-9A-Za-z_-]+)>",
        r"<([0-9A-Za-z_]+)>"
    ]

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    doc = word.Documents.Open(doc_path, ReadOnly=False)

    try:
        # Loop all stories
        story_range = doc.StoryRanges

        for rng in story_range:
            current = rng
            while current is not None:
                text = current.Text
                original = text

                # 1️⃣ Remove all <tags>
                for pat in tag_patterns:
                    text = re.sub(pat, "", text)

                # 2️⃣ Remove leading spaces in paragraphs
                text = re.sub(r"(?m)^\s+", "", text)

                # 3️⃣ Replace multiple spaces with single
                text = re.sub(r" {2,}", " ", text)

                # 4️⃣ Remove blank lines (empty paragraph markers)
                text = re.sub(r"(?m)^\s*$\r?", "", text)

                # Only update if changed
                if text != original:
                    current.Text = text

                current = current.NextStoryRange

        # Save the SAME file
        doc.Save()

    finally:
        doc.Close(SaveChanges=False)
        word.Quit()
        pythoncom.CoUninitialize()

    return doc_path




# --- keep your remaining formatting, multilingual, and HTML helper functions here ---
# (from your working version)

import pythoncom
import win32com.client
import re

def highlight_keywords_plus_next_word_com(doc):
    """
    Highlights keywords + next word in Word document using COM automation.
    Returns:
        True  -> at least one highlight applied
        False -> no highlight applied
    """
    keywords = [
        "Refer", "Insert", "Pick-up", "pickup", "See",
        "COMP", "AU", "AQ", "SPU", "Compositor",
        "Ph", "Photo", "video", "images"
    ]

    highlight_done = False
    pattern = r'\b(' + '|'.join(re.escape(k) for k in keywords) + r')\b\s+(\S+)'

    # Iterate paragraphs
    for para in doc.Paragraphs:
        rng = para.Range
        text = rng.Text

        for match in re.finditer(pattern, text, flags=re.IGNORECASE):
            start, end = match.span()
            match_range = rng.Duplicate
            match_range.Start = rng.Start + start
            match_range.End = rng.Start + end

            try:
                match_range.HighlightColorIndex = 4  # wdYellow
            except:
                match_range.Font.HighlightColorIndex = 4

            highlight_done = True   # <<-- FLAG SET

    return highlight_done

# ------------------------------
# Document extraction helpers
# ------------------------------
def extract_with_word(doc_path: str):
    """
    Uses Word automation (pywin32).
    Returns:
        paragraphs, comments, img_count, footnotes, endnotes
    Also:
        Saves document ONLY IF highlights were applied.
    """
    pythoncom.CoInitialize()
    word = None

    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(doc_path, ReadOnly=False)

        try:
            doc.Repaginate()
        except:
            pass

        analyzer = CitationAnalyzer()

        # >>> Detect keyword highlighting
        keyword_highlighted = highlight_keywords_plus_next_word_com(doc)

        paragraphs = []
        for para in doc.Paragraphs:
            txt = para.Range.Text.strip('\r\x07')
            if not txt:
                continue
            page_no = para.Range.Information(3)

            is_highlighted = (para.Range.HighlightColorIndex != 0)
            is_caption = analyzer.is_caption_paragraph(txt)

            paragraphs.append((txt, page_no, is_caption, is_highlighted))

        comments = []
        for c in doc.Comments:
            try:
                comments.append((c.Author, c.Range.Text.strip('\r'), c.Scope.Information(3)))
            except:
                continue

        img_count = doc.InlineShapes.Count + sum(1 for s in doc.Shapes if s.Type in (13, 11))
        footnotes = doc.Footnotes.Count
        endnotes = doc.Endnotes.Count

        # >>> SAVE ONLY IF HIGHLIGHTS WERE APPLIED
        if keyword_highlighted:
            doc.Save()

        doc.Close(SaveChanges=False)
        word.Quit()

        return paragraphs, comments, img_count, footnotes, endnotes

    except Exception as e:
        raise Exception(f"Word extraction failed: {e}")

    finally:
        if word:
            try:
                word.Quit()
            except:
                pass
        pythoncom.CoUninitialize()




def extract_with_docx(doc_path: str):
    """
    Fallback extraction with python-docx.
    Returns:
    - paragraphs: list of (text, approx_page_no, is_caption, is_highlighted=False)
    - comments: [] (python-docx doesn't surface comments reliably)
    - image_count: approximate inline_shapes count
    - footnote_count, endnote_count: 0 (not supported by python-docx)
    """
    doc = Document(doc_path)
    paragraphs = []
    analyzer = CitationAnalyzer()
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        # page number unknown — approximate as paragraph index (1-based)
        approx_page = i // 40 + 1  # crude heuristic: 40 paras ~= 1 page (very rough)
        is_caption = analyzer.is_caption_paragraph(text)
        # python-docx doesn't give highlight info at API level
        paragraphs.append((text, f"approx.{approx_page}", is_caption, False))

    # image count: look for shapes in document._part (count pictures)
    imgs = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            imgs += 1

    return paragraphs, [], imgs, 0, 0


# ------------------------------
# HTML Template pieces
# ------------------------------
# CSS + JS embedded (adapted from your VBA)
DASHBOARD_CSS = r"""/* === S4Carlisle AI Manuscript Analysis Dashboard === */
* { margin: 0; padding: 0; box-sizing: border-box; }
body { font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); min-height: 100vh; color: #333; padding: 20px; }
.container { max-width: 1400px; margin: 0 auto; }
.header { background: white; border-radius: 15px; padding: 30px; margin-bottom: 20px; box-shadow: 0 10px 30px rgba(0,0,0,0.1); }
.title { font-size: 2rem; font-weight: 700; color: #2c3e50; margin-bottom: 20px; }
.metadata { display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap: 15px; }
.meta-item { background: #f8f9ff; padding: 15px; border-radius: 10px; border-left: 4px solid #667eea; }
.meta-label { font-weight: 600; color: #555; font-size: 0.9rem; }
.meta-value { font-size: 1.1rem; font-weight: 700; color: #2c3e50; margin-top: 5px; }
.nav-tabs { display: flex; background: white; border-radius: 15px; padding: 5px; margin-bottom: 20px; box-shadow: 0 5px 15px rgba(0,0,0,0.1); gap: 5px; }
.nav-tab { flex: 1; text-align: center; padding: 15px; border-radius: 10px; cursor: pointer; transition: all 0.3s; font-weight: 500; }
.nav-tab.active { background: #667eea; color: white; box-shadow: 0 5px 15px rgba(102, 126, 234, 0.3); }
.nav-tab:hover:not(.active) { background: #f8f9ff; }
.tab-content { display: none; background: white; border-radius: 15px; padding: 30px; box-shadow: 0 8px 25px rgba(0,0,0,0.1); }
.tab-content.active { display: block; animation: fadeIn 0.3s; }
@keyframes fadeIn { from { opacity: 0; } to { opacity: 1; } }
.section-title { font-size: 1.5rem; font-weight: 600; color: #2c3e50; margin-bottom: 20px; }
table { width: 100%; border-collapse: collapse; margin: 20px 0; background: white; border-radius: 10px; overflow: hidden; }
th { background: #667eea; color: white; padding: 12px; font-weight: 600; text-align: left; }
td { padding: 10px 12px; border-bottom: 1px solid #eee; }
tr:hover { background: #f8f9ff; }
h3 { color: #2c3e50; margin-top: 30px; margin-bottom: 15px; font-size: 1.2rem; }
.summary-table { margin-bottom: 30px; }
@media (max-width: 768px) { .container { padding: 10px; } .title { font-size: 1.5rem; } .metadata { grid-template-columns: 1fr; } }
"""

DASHBOARD_JS = r"""
function showTab(tabId) {
    document.querySelectorAll('.tab-content').forEach(c => c.classList.remove('active'));
    document.querySelectorAll('.nav-tab').forEach(t => t.classList.remove('active'));
    const target = document.getElementById(tabId);
    if (target) target.classList.add('active');
    // set active on the tab that was clicked
    const tabs = document.querySelectorAll('.nav-tab');
    tabs.forEach(tab => {
        if (tab.getAttribute('data-target') === tabId) tab.classList.add('active');
    });
}
document.addEventListener('DOMContentLoaded', function() {
    // make first tab active if not already
    if (!document.querySelector('.nav-tab.active')) {
        const first = document.querySelector('.nav-tab');
        if (first) first.classList.add('active');
    }
});
"""

# Basic HTML page wrapper (Jinja-like placeholders)
HTML_WRAPPER = """<!doctype html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>Document Analysis - {{ doc_name }}</title>
<link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.0.0/css/all.min.css">
<link rel="stylesheet" href="https://cdn.datatables.net/1.13.6/css/jquery.dataTables.min.css">
<style>{{ css }}</style>
<script src="https://code.jquery.com/jquery-3.7.1.min.js"></script>
<script src="https://cdn.datatables.net/1.13.6/js/jquery.dataTables.min.js"></script>
</head>
<body>
<div class="container">
    <div class="header">
        <div class="title">
            <img src="{{ logo_path }}" alt="" style="height:40px;vertical-align:middle;margin-right:10px;">
            <i class="fa-solid fa-robot"></i>S4Carlisle Manuscript Analysis Dashboard
        </div>
        <div class="metadata">
            <div class="meta-item"><div class="meta-label">File</div><div class="meta-value">{{ doc_name }}</div></div>
            <div class="meta-item"><div class="meta-label">Pages</div><div class="meta-value">{{ pages }}</div></div>
            <div class="meta-item"><div class="meta-label">Words</div><div class="meta-value">{{ words }}</div></div>
            <div class="meta-item"><div class="meta-label">CE Pages</div><div class="meta-value">{{ ce_pages }}</div></div>
            <div class="meta-item"><div class="meta-label">Date</div><div class="meta-value">{{ date }}</div></div>
            <div class="meta-item"><div class="meta-label">Analyst</div><div class="meta-value">{{ analyst }}</div></div>
        </div>
    </div>
    
    <div id="analysis-summary" class="tab-content active" style="margin-bottom: 25px;">
    {{ detailed_summary|safe }}
    </div>

    <!-- Navigation Tabs -->
    <div class="nav-tabs">
        <div class="nav-tab active" data-target="citations" onclick="showTab('citations')">Citations</div>
        <div class="nav-tab" data-target="special-chars" onclick="showTab('special-chars')">Special Chars</div>
        <div class="nav-tab" data-target="formatting" onclick="showTab('formatting')">Formatting</div>
        <div class="nav-tab" data-target="comments" onclick="showTab('comments')">Comments</div>
        <div class="nav-tab" data-target="media" onclick="showTab('media')">Media</div>
    </div>

    <!-- Tabs -->
    <div id="citations" class="tab-content active">
        <div class="section-title"><i class="fa-solid fa-closed-captioning"></i> Citations & Captions</div>
        {{ msr_content|safe }}
    </div>

    <div id="special-chars" class="tab-content">
        <div class="section-title"><i class="fas fa-language"></i> Special Characters</div>
        {{ spec_content|safe }}
    </div>

    <div id="formatting" class="tab-content">
        <div class="section-title"><i class="fas fa-cogs"></i> Formatting</div>
        {{ fmt_content|safe }}
    </div>

    <div id="comments" class="tab-content">
        <div class="section-title"><i class="fas fa-comments"></i> Comments & Highlights</div>
        {{ comment_content|safe }}
        {{ export_highlight|safe }}
    </div>

    <div id="media" class="tab-content">
        <div class="section-title"><i class="fas fa-images"></i> Media & Notes</div>
        <p><b>Images:</b> {{ images }} | <b>Footnotes:</b> {{ footnotes }} | <b>Endnotes:</b> {{ endnotes }}</p>
    </div>

</div>

<!-- Tab JS -->
<script>
function showTab(tabId) {
    document.querySelectorAll('.tab-content').forEach(c => c.classList.remove('active'));
    document.querySelectorAll('.nav-tab').forEach(t => t.classList.remove('active'));
    const target = document.getElementById(tabId);
    if (target) target.classList.add('active');
    document.querySelectorAll('.nav-tab').forEach(tab => {
        if (tab.getAttribute('data-target') === tabId) tab.classList.add('active');
    });
}
</script>

<script>
$(document).ready(function(){
    $('table').each(function(){
        // Skip tables with irregular rows (colspan/rowspan)
        const hasIrregularRows = $(this).find('td[colspan], td[rowspan]').length > 0;
        if (hasIrregularRows) {
            console.log('Skipping DataTables init for irregular table:', this.id);
            return; // ✅ Skip DataTables for this table
        }

        try {
            $(this).DataTable({
                pageLength: 10,
                autoWidth: false,
                ordering: true,
                responsive: true,
                columnDefs: [
                    { targets: "_all", defaultContent: "" }
                ]
            });
        } catch (e) {
            console.warn('DataTable init failed for', this.id, e);
        }
    });
});
</script>


<script>{{ js }}</script>
</body>
</html>
"""
# ------------------------------
# Helper pieces ported from VBA (best-effort)
# ------------------------------
def generate_formatting_html(doc_path: str, used_word: bool) -> str:
    # Best-effort formatting scan
    rows = []
    if used_word and HAS_WIN32COM:
        # use Word to detect strikethrough / hidden / text boxes & section breaks
        pythoncom.CoInitialize()
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(doc_path, ReadOnly=True)
        try:
            rng = doc.Content
            # Strikethrough
            rng.Find.ClearFormatting()
            rng.Find.Font.StrikeThrough = True
            rng.Find.Text = ""
            rng.Find.Forward = True
            rng.Find.Format = True
            while rng.Find.Execute():
                page = rng.Information(3)
                rows.append(("Formatting", page, "Strikethrough", escape_html(rng.Text.strip())))
                rng.Collapse(0)
            # Hidden
            rng = doc.Content
            rng.Find.ClearFormatting()
            rng.Find.Font.Hidden = True
            rng.Find.Text = ""
            rng.Find.Forward = True
            rng.Find.Format = True
            while rng.Find.Execute():
                page = rng.Information(3)
                rows.append(("Formatting", page, "Hidden", escape_html(rng.Text.strip())))
                rng.Collapse(0)
            # Section breaks
            for sec in doc.Sections:
                rows.append(("Formatting", sec.Range.Information(3), "Section Break", "(Section Break)"))
            # Text frames: doc.Shapes
            for shp in doc.Shapes:
                try:
                    if shp.Type == 17:  # msoTextBox sometimes varies; fallback to reading text
                        anchor_page = shp.Anchor.Information(3)
                        rows.append(("Formatting", anchor_page, "Text Frame", escape_html(shp.TextFrame.TextRange.Text.strip())))
                except Exception:
                    pass
        finally:
            doc.Close(False)
            word.Quit()
    else:
        # fallback: quick heuristics using python-docx runs to find strikethrough or hidden (python-docx doesn't expose hidden)
        if HAS_DOCX:
            doc = Document(doc_path)
            for i, p in enumerate(doc.paragraphs):
                for run in p.runs:
                    if getattr(run.font, "strike", False):
                        rows.append(("Formatting", f"approx.{i//40+1}", "Strikethrough", escape_html(run.text)))
            # section breaks: approximate by page breaks
            rows.append(("Formatting", "N/A", "Note", "Nil"))
        else:
            rows.append(("Formatting", "N/A", "Note", "Nil"))

    # Build table HTML
    html = "<table><thead><tr><th>Type</th><th>Page</th><th>Category</th><th>Details</th></tr></thead><tbody>"
    if rows:
        for r in rows:
            html += f"<tr><td>{r[0]}</td><td>{r[1]}</td><td>{r[2]}</td><td>{r[3]}</td></tr>"
    else:
        html += "<tr><td colspan='4'>No formatting issues found or not detectable without Word automation.</td></tr>"
    html += "</tbody></table>"
    return html


from collections import defaultdict
import pythoncom
import win32com.client
from docx import Document

from collections import defaultdict
import pythoncom

# Optional imports
try:
    import win32com.client
    HAS_WIN32COM = True
except ImportError:
    HAS_WIN32COM = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def page_ranges(pages):
    """Convert sorted set/list of ints to compact ranges: 1-3,5,7-9"""
    if not pages:
        return ""
    pages = sorted(pages)
    ranges = []
    start = prev = pages[0]
    for p in pages[1:]:
        if p == prev + 1:
            prev = p
        else:
            ranges.append(f"{start}" if start == prev else f"{start}-{prev}")
            start = prev = p
    ranges.append(f"{start}" if start == prev else f"{start}-{prev}")
    return ", ".join(ranges)


from collections import defaultdict
import pythoncom
import win32com.client

def generate_multilingual_html(doc_path: str) -> str:
    """
    Uses merged highlighter. Highlights directly in same file.
    Returns HTML for multilingual characters only.
    """
    pythoncom.CoInitialize()
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False

    doc = word.Documents.Open(doc_path, ReadOnly=False)

    try:
        try:
            doc.Repaginate()
        except:
            pass

        # 🔥 Our new merged highlighter
        page_map, highlighted = highlight_all_in_one(doc)

        # Save only if changed
        if highlighted:
            doc.Save()

    finally:
        doc.Close(SaveChanges=False)
        word.Quit()
        pythoncom.CoUninitialize()

    # Build HTML summary
    html = "<table><thead><tr><th>Language/Type</th><th>Page</th></tr></thead><tbody>"
    for lang, pages in page_map.items():
        for p in sorted(pages):
            html += f"<tr><td>{lang}</td><td>{p}</td></tr>"

    if not page_map:
        html += "<tr><td colspan='2'>No multilingual characters found</td></tr>"

    html += "</tbody></table>"
    return html


def highlight_all_in_one(doc):
    """
    Performs ALL highlighting in ONE pass:
    - Keyword next-word highlight
    - Multilingual character highlight
    Returns:
        page_map: dict of multilingual pages
        highlighted: True/False if ANY highlight was applied
    """

    keywords = [
        "Refer", "Insert", "Pick-up", "pickup", "See",
        "COMP", "AU", "AQ", "SPU", "Compositor",
        "Ph", "Photo", "video", "images"
    ]

    keyword_pattern = r'\b(' + '|'.join(re.escape(k) for k in keywords) + r')\b\s+(\S+)'

    multilingual_ranges = [
        ("Chinese",      0x4E00, 0x9FFF),
        ("Greek",        0x0370, 0x03FF),
        ("Cyrillic",     0x0400, 0x04FF),
        ("Hebrew",       0x0590, 0x05FF),
        ("Arabic",       0x0600, 0x06FF),
        ("Arabic",       0x0750, 0x077F),
        ("Devanagari",   0x0900, 0x097F),
        ("Japanese",     0x3040, 0x309F),
        ("Japanese",     0x30A0, 0x30FF),
        ("Korean",       0xAC00, 0xD7AF),
        ("Thai",         0x0E00, 0x0E7F),
        ("Currency",     0x20A0, 0x20CF),
    ]

    page_map = defaultdict(set)
    highlighted = False

    total_pages = doc.ComputeStatistics(2)

    # SINGLE COM LOOP — MUCH FASTER
    for para in doc.Paragraphs:
        rng = para.Range
        text = rng.Text
        page_no = rng.Information(3)
        if page_no > total_pages:
            page_no = total_pages

        # -----------------------------------
        # 1) Keyword highlight (keyword + next word)
        # -----------------------------------
        for match in re.finditer(keyword_pattern, text, flags=re.IGNORECASE):
            start, end = match.span()
            r = rng.Duplicate
            r.Start = rng.Start + start
            r.End   = rng.Start + end

            try:
                r.HighlightColorIndex = 7    # Yellow
            except:
                r.Font.HighlightColorIndex = 7

            highlighted = True

        # -----------------------------------
        # 2) Multilingual character detection
        # -----------------------------------
        for i, ch in enumerate(text):
            code = ord(ch)

            for lang, low, high in multilingual_ranges:
                if low <= code <= high:

                    # Highlight character
                    char_r = rng.Duplicate
                    char_r.Start = rng.Start + i
                    char_r.End   = rng.Start + i + 1

                    try:
                        char_r.HighlightColorIndex = 4   # BrightGreen
                    except:
                        char_r.Font.HighlightColorIndex = 4

                    highlighted = True

                    # Add to HTML table
                    page_map[lang].add(page_no)
                    break

    return page_map, highlighted






def build_comments_html(comments: List[Tuple]):
    if not comments:
        return "<p>No comments found or comments unavailable (python-docx can't always read comments).</p>"
    html = "<table><thead><tr><th>#</th><th>Page</th><th>Author</th><th>Comment</th></tr></thead><tbody>"
    for i, (author, text, page) in enumerate(comments, start=1):
        html += f"<tr><td>{i}</td><td>{page}</td><td>{escape_html(author)}</td><td>{escape_html(text)}</td></tr>"
    html += "</tbody></table>"
    return html


def build_export_highlight_html(paragraphs_full):
    # paragraphs_full elements: (text, page, is_caption, is_highlighted) if word used
    highlights = []
    for t, p, is_cap, is_high in paragraphs_full:
        if is_high:
            highlights.append((t, p))
    if not highlights:
        return "<p>No highlighted paragraphs found.</p>"
    html = "<table><thead><tr><th>Highlighted Text</th><th>Page</th></tr></thead><tbody>"
    for t, p in highlights:
        html += f"<tr><td>{escape_html(t)}</td><td>{p}</td></tr>"
    html += "</tbody></table>"
    return html


def escape_html(s: str) -> str:
    return (s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")
            .replace("\n", "<br>"))


__all__ = [
    "CitationAnalyzer",
    "extract_with_word",
    "extract_with_docx",
    "generate_formatting_html",
    "generate_multilingual_html",
    "build_comments_html",
    "build_export_highlight_html",
    "DASHBOARD_CSS",
    "DASHBOARD_JS",
    "HTML_WRAPPER",
]
